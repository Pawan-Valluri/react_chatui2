# backend/agent.py

import contextvars
import os
import docx
from typing import Annotated, List, TypedDict, Any, Literal, Optional
from langchain_core.tools import tool
from langchain_core.language_models import BaseChatModel
from langchain_core.messages import BaseMessage, AIMessage, HumanMessage, ToolMessage
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode
from app.services.document_service import current_thread_id_var

current_document_context_var = contextvars.ContextVar("current_document_context")

def get_context_document_text() -> str:
    try:
        ctx = current_document_context_var.get()
        if ctx and "content" in ctx:
            content = ctx["content"]
            lines = []
            for block in content:
                b_type = block.get("type")
                if b_type == "paragraph":
                    para_id = block.get("paraId")
                    text = block.get("text", "")
                    id_str = f"[{para_id}] " if para_id else ""
                    lines.append(f"{id_str}{text}")
                elif b_type == "table":
                    rows = block.get("rows", [])
                    for r_idx, row in enumerate(rows):
                        cells = row.get("cells", [])
                        for c_idx, cell in enumerate(cells):
                            cell_content = cell.get("content", [])
                            for cell_block in cell_content:
                                if cell_block.get("type") == "paragraph":
                                    para_id = cell_block.get("paraId")
                                    text = cell_block.get("text", "")
                                    id_str = f"[{para_id}] " if para_id else ""
                                    lines.append(f"{id_str}(table, row {r_idx + 1}, col {c_idx + 1}) {text}")
            return "\n".join(lines)
    except LookupError:
        pass
    
    # Fallback to reading file on disk
    try:
        from app.services.document_service import get_document_path
        thread_id = current_thread_id_var.get()
        path = get_document_path(thread_id)
        if os.path.exists(path):
            doc = docx.Document(path)
            lines = []
            for p in doc.paragraphs:
                w14_ns = "http://schemas.microsoft.com/office/word/2010/wordml"
                para_id = p._element.get(f"{{{w14_ns}}}paraId")
                if not para_id:
                    para_id = p._element.get("paraId")
                id_str = f"[{para_id}] " if para_id else ""
                lines.append(f"{id_str}{p.text}")
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        for p in cell.paragraphs:
                            w14_ns = "http://schemas.microsoft.com/office/word/2010/wordml"
                            para_id = p._element.get(f"{{{w14_ns}}}paraId")
                            if not para_id:
                                para_id = p._element.get("paraId")
                            id_str = f"[{para_id}] " if para_id else ""
                            lines.append(f"{id_str}(table, row {r_idx + 1}, col {c_idx + 1}) {p.text}")
            return "\n".join(lines)
    except Exception as e:
        print("Fallback read failed:", e)
    return "Error: Document content not available."

# Define clean enterprise Agent State structure
class AgentState(TypedDict):
    messages: Annotated[List[BaseMessage], add_messages]
    reasoning_steps: List[str]  # Tracks the consecutive reasoning thoughts
    loop_counter: int           # Safety guardrail counter

# ──────────────────────────────────────────────────────────────────────────
# 🛠️ Define Clean Native Production Tools
# ──────────────────────────────────────────────────────────────────────────

@tool
def think(thought: str) -> str:
    """Use this tool to record consecutive thoughts, plans, or silent reasoning steps."""
    return "Thought recorded."

@tool
async def search_kb(query: str) -> str:
    """Search the corporate Knowledge Base for APCOT Chat system specifications, vanilla CSS design principles, and guidelines."""
    return (
        "Knowledge Base Found: 'APCOT Chat' is a premium web client utilizing "
        "@assistant-ui/react primitives styled completely via Vanilla CSS (Tailwind-free) "
        "complying with modular component-tree encapsulation design principles."
    )

@tool
async def check_entitlements(resource: str) -> str:
    """Check group entitlements and access permissions for the current resource."""
    return (
        "AUTHORIZED: User is a member of 'SSO_APP_ADMIN'. "
        "Granted full administration, thread deletion, and query permissions."
    )

@tool
def read_document(fromIndex: Optional[int] = None, toIndex: Optional[int] = None) -> str:
    """Read the document content. Returns lines tagged with a stable paragraph id, e.g. '[2A1F3B] First paragraph'."""
    text = get_context_document_text()
    if not text:
        return "Empty document."
    lines = text.split("\n")
    start = fromIndex if fromIndex is not None else 0
    end = (toIndex + 1) if toIndex is not None else len(lines)
    return "\n".join(lines[start:end])

@tool
def read_selection() -> str:
    """Read the user's current cursor or selection in the editor."""
    try:
        ctx = current_document_context_var.get()
        if ctx and "selection" in ctx and ctx["selection"]:
            sel = ctx["selection"]
            return f"Selection: Text '{sel.get('selectedText', '')}' in paragraph {sel.get('paraId', '')}."
    except LookupError:
        pass
    return "No active user selection."

@tool
def read_page(pageNumber: int) -> str:
    """Read the contents of one rendered page (1-indexed)."""
    text = get_context_document_text()
    lines = text.split("\n")
    if pageNumber <= 1:
        return "\n".join(lines[:15]) if len(lines) > 15 else text
    else:
        return "Page not found."

@tool
def read_pages(from_page: int, to_page: int) -> str:
    """Read a contiguous range of rendered pages (1-indexed, inclusive)."""
    text = get_context_document_text()
    return f"--- Pages {from_page} to {to_page} ---\n{text}"

@tool
def find_text(query: str, caseSensitive: bool = False, limit: int = 20) -> str:
    """Locate paragraphs containing query. Returns paraId and matches."""
    text = get_context_document_text()
    lines = text.split("\n")
    results = []
    q = query if caseSensitive else query.lower()
    for line in lines:
        match_line = line if caseSensitive else line.lower()
        if q in match_line:
            results.append(line)
            if len(results) >= limit:
                break
    return "\n".join(results) if results else "No matches found."

@tool
def read_comments() -> str:
    """List all comments in the document with their paragraph anchors."""
    try:
        ctx = current_document_context_var.get()
        if ctx and "comments" in ctx and ctx["comments"]:
            comments = ctx["comments"]
            res = []
            for c in comments:
                replies = ", ".join([f"{r.get('author')}: {r.get('text')}" for r in c.get('replies', [])])
                replies_str = f" (Replies: {replies})" if replies else ""
                res.append(f"[Comment #{c.get('id')}] {c.get('author')}: \"{c.get('text')}\" on {c.get('paraId')}{replies_str}")
            return "\n".join(res)
    except LookupError:
        pass
    return "No comments in document."

@tool
def read_changes() -> str:
    """List tracked changes (insertions / deletions) currently in the document."""
    try:
        ctx = current_document_context_var.get()
        if ctx and "changes" in ctx and ctx["changes"]:
            changes = ctx["changes"]
            res = []
            for c in changes:
                res.append(f"[Change #{c.get('id')}] {c.get('type')} by {c.get('author')}: \"{c.get('text')}\"")
            return "\n".join(res)
    except LookupError:
        pass
    return "No tracked changes in document."

@tool
def add_comment(paraId: str, text: str, search: Optional[str] = None) -> str:
    """Attach a comment to a paragraph, optionally anchored to a unique phrase within it."""
    return f"Comment addition queued for paraId: {paraId}."

@tool
def suggest_change(paraId: str, search: str, replaceWith: str) -> str:
    """Suggest a tracked change (replacement, deletion, or insertion)."""
    return f"Change suggestion queued for paraId: {paraId}."

@tool
def apply_formatting(paraId: str, marks: dict, search: Optional[str] = None) -> str:
    """Apply character formatting (bold, italic, underline, strike, color, highlight, fontSize, fontFamily) to a paragraph or to a unique phrase within it."""
    return f"Formatting application queued for paraId: {paraId}."

@tool
def set_paragraph_style(paraId: str, styleId: str) -> str:
    """Apply a paragraph style by id (e.g. 'Heading1', 'Title', 'Quote', 'Normal')."""
    return f"Style '{styleId}' application queued for paraId: {paraId}."

@tool
def reply_comment(commentId: int, text: str) -> str:
    """Reply to an existing comment by id."""
    return f"Reply queued for comment ID: {commentId}."

@tool
def resolve_comment(commentId: int) -> str:
    """Mark a comment as resolved (done)."""
    return f"Resolution queued for comment ID: {commentId}."

@tool
def scroll(paraId: str) -> str:
    """Scroll the editor to a paragraph by paraId."""
    return f"Viewport scroll queued for paraId: {paraId}."

@tool
def insert_table(paraId: str, rows: int, cols: int, styleId: Optional[str] = None, cells: Optional[List[List[dict]]] = None) -> str:
    """Insert a highly customizable table structure after the paragraph specified by paraId.
    - cells: Optional 2D list containing cell configuration objects.
      Each cell object can have:
        - 'shading': Hex color string (e.g., 'E6F0FA').
        - 'content': List of block elements inside the cell.
          Block elements can be:
            - {'type': 'paragraph', 'text': '...', 'marks': {'bold': True}}
            - {'type': 'bullet_list', 'items': ['Item 1', 'Item 2']}
    """
    return f"Table insertion queued after paraId: {paraId}."

@tool
def edit_table_cell(paraId: str, cellIndex: Optional[dict] = None, shading: Optional[str] = None, content: Optional[List[dict]] = None) -> str:
    """Update cell content (paragraphs, bullet lists) and background shading for an existing table cell."""
    return f"Table cell update queued for cell containing paraId: {paraId}."

@tool
def toggle_bullet_list(paraId: str, enabled: bool) -> str:
    """Convert a paragraph with paraId into a bullet point list item, or revert it to a normal paragraph."""
    return f"Bullet list toggle queued for paraId: {paraId}."

@tool
def append_paragraph(text: str, styleId: str = "Normal") -> str:
    """Append a new paragraph to the end of the document."""
    return f"Paragraph appendage queued with text: {text[:20]}..."

@tool
def add_table_row(paraId: str, position: str = "after") -> str:
    """Insert a new table row relative to the row containing paraId. position can be 'before' or 'after'."""
    return f"Table row insertion queued relative to paraId: {paraId}."

@tool
def delete_table_row(paraId: str) -> str:
    """Delete the table row containing paraId."""
    return f"Table row deletion queued for row containing paraId: {paraId}."

@tool
def add_table_column(paraId: str, position: str = "after") -> str:
    """Insert a new table column relative to the column containing paraId. position can be 'before' or 'after'."""
    return f"Table column insertion queued relative to paraId: {paraId}."

@tool
def delete_table_column(paraId: str) -> str:
    """Delete the table column containing paraId."""
    return f"Table column deletion queued for column containing paraId: {paraId}."

# ──────────────────────────────────────────────────────────────────────────
# 📚 Inject Markdown Style Repositories into Tool Descriptions
# ──────────────────────────────────────────────────────────────────────────
import os

ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "styles")

try:
    with open(os.path.join(ASSETS_DIR, "paragraph_styles.md"), "r") as f:
        PARAGRAPH_STYLES_MD = f.read()
    set_paragraph_style.description += f"\n\n{PARAGRAPH_STYLES_MD}"
    append_paragraph.description += f"\n\n{PARAGRAPH_STYLES_MD}"
except Exception as e:
    print("Warning: Could not load paragraph styles markdown.", e)

try:
    with open(os.path.join(ASSETS_DIR, "table_styles.md"), "r") as f:
        TABLE_STYLES_MD = f.read()
    insert_table.description += f"\n\n{TABLE_STYLES_MD}"
except Exception as e:
    print("Warning: Could not load table styles markdown.", e)

# ──────────────────────────────────────────────────────────────────────────
# 🔀 The Dynamic Conditional Router
# ──────────────────────────────────────────────────────────────────────────

def evaluate_agent_step(state: AgentState) -> Literal["agent", "tools", "__end__"]:
    # Loop safety guardrail to prevent infinite billing recursion
    if state.get("loop_counter", 0) >= 25:
        print("Loop Safety Guardrail triggered: forcing termination.")
        return END

    last_message = state["messages"][-1]

    # If the last message contains no tool calls, we have our final response!
    if not last_message.tool_calls:
        return END

    # Check if the LLM called the virtual 'think' tool
    is_thinking = any(tc["name"] == "think" for tc in last_message.tool_calls)
    if is_thinking:
        return "agent"  # Loops back to agent_node for consecutive thought steps

    return "tools"      # Routes to native ToolNode for system execution

# ──────────────────────────────────────────────────────────────────────────
# 🏗️ State Machine Factory
# ──────────────────────────────────────────────────────────────────────────

def create_agent_graph(llm: BaseChatModel, tools: List[Any] = None):
    """Compiles the pristine production LangGraph using a polymorphic LLM dependency."""
    
    # 1. Default to core native production tools if none are passed
    if tools is None:
        tools = [
            think, search_kb, check_entitlements,
            read_document, read_selection, read_page, read_pages,
            find_text, read_comments, read_changes,
            add_comment, suggest_change, apply_formatting,
            set_paragraph_style, reply_comment, resolve_comment, scroll,
            insert_table, edit_table_cell, toggle_bullet_list,
            add_table_row, delete_table_row, add_table_column, delete_table_column,
            append_paragraph
        ]
        
    llm_with_tools = llm.bind_tools(tools)

    # 2. Define the Agent thinking node
    async def agent_node(state: AgentState) -> dict:
        # Log previous messages to console to test correct branching
        print(f"\n--- [AGENT GRAPH RUN] Dynamic Message Context (Count: {len(state['messages'])}) ---")
        for idx, msg in enumerate(state["messages"]):
            role = "USER" if msg.type == "human" else ("TOOL" if msg.type == "tool" else "ASSISTANT")
            details = ""
            if msg.type == "ai" and msg.tool_calls:
                details = f" [ToolCalls: {[tc['name'] for tc in msg.tool_calls]}]"
            elif msg.type == "tool":
                details = f" [ToolName: {msg.name}, Status: complete]"
            snippet = msg.content[:90] if isinstance(msg.content, str) else str(msg.content)[:90]
            print(f"  [{idx}] {role}: {snippet}{details}")
        print("--------------------------------------------------------------------\n")

        # Invoke LLM
        response = await llm_with_tools.ainvoke(state["messages"])
        
        # Extract and compile reasoning steps
        new_reasoning = list(state.get("reasoning_steps", []))
        if response.tool_calls:
            for tc in response.tool_calls:
                if tc["name"] == "think":
                    new_reasoning.append(tc["args"]["thought"])

        return {
            "messages": [response],
            "reasoning_steps": new_reasoning,
            "loop_counter": state.get("loop_counter", 0) + 1
        }

    # 3. Compile workflow
    workflow = StateGraph(AgentState)
    
    # Add nodes
    workflow.add_node("agent", agent_node)
    workflow.add_node("tools", ToolNode(tools))
    
    # Add edges
    workflow.add_edge(START, "agent")
    workflow.add_conditional_edges("agent", evaluate_agent_step, ["agent", "tools", END])
    workflow.add_edge("tools", "agent")
    
    return workflow.compile()
