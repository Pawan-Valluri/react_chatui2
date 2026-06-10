# backend/mock_showcase/mock_llm.py

import json
import uuid
import time
import os
from typing import List, Optional, Any, Dict, Sequence
from langchain_core.language_models import BaseChatModel
from langchain_core.messages import BaseMessage, AIMessage, HumanMessage, ToolMessage
from langchain_core.outputs import ChatResult, ChatGeneration
from langchain_core.callbacks import CallbackManagerForLLMRun

from app.mock_showcase.mock_routes import CONVERSATION_ROUTES

# Helper to load config.json from root directory
def load_combined_config():
    config = {
        "AGENT_THINKING_DELAY": 0.1,
        "AGENT_TOOL_TRIGGER_DELAY": 0.1,
        "AGENT_TOOL_EXECUTION_DELAY": 0.2
    }
    try:
        config_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "config.json")
        if os.path.exists(config_path):
            with open(config_path, "r", encoding="utf-8") as f:
                loaded = json.load(f)
                mock_cfg = loaded.get("MOCK_CONFIG", {})
                for k, v in mock_cfg.items():
                    config[k] = v
    except Exception as e:
        print("Failed to load root config.json in mock_llm:", e)
    return config

def get_para_ids() -> List[str]:
    para_ids = []
    try:
        from app.agent import current_document_context_var
        ctx = current_document_context_var.get()
        if ctx and "content" in ctx:
            for block in ctx["content"]:
                if block.get("type") == "paragraph":
                    p_id = block.get("paraId")
                    if p_id:
                        para_ids.append(p_id)
    except Exception as e:
        print("Context read failed in mock_llm:", e)
        
    if para_ids:
        return para_ids

    # Fallback to reading file on disk
    try:
        from app.services.document_service import get_document_path, current_thread_id_var
        import docx
        thread_id = current_thread_id_var.get()
        path = get_document_path(thread_id)
        if os.path.exists(path):
            doc = docx.Document(path)
            for p in doc.paragraphs:
                w14_ns = "http://schemas.microsoft.com/office/word/2010/wordml"
                para_id = p._element.get(f"{{{w14_ns}}}paraId")
                if not para_id:
                    para_id = p._element.get("paraId")
                if para_id:
                    para_ids.append(para_id)
    except Exception as e:
        print("Fallback file read failed in mock_llm:", e)
        
    if not para_ids:
        para_ids = ["p_first", "p_second", "p_third"]
    return para_ids

class MockChatModel(BaseChatModel):
    """A polymorphic LLM simulator that acts exactly like a real LangChain model,
    implementing state-based predefined showcase pathways.
    """
    
    # Track current user info for dynamic template interpolation
    user_id: str = "beyond_dev"
    fullname: str = "Beyond Developer"

    @property
    def _llm_type(self) -> str:
        return "mock-showcase-chat"

    def bind_tools(self, tools: Sequence[Any], **kwargs: Any) -> "MockChatModel":
        """Polymorphic override of bind_tools to prevent NotImplementedError."""
        return self

    def _generate(
        self,
        messages: List[BaseMessage],
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> ChatResult:
        # Simulate thinking delay in mock showcase
        mock_cfg = load_combined_config()
        thinking_delay = float(mock_cfg.get("AGENT_THINKING_DELAY", 0.1))
        if thinking_delay > 0:
            time.sleep(thinking_delay)

        # Log previous messages in the LLM call to verify correctness of branching
        print(f"\n================ [MOCK LLM INVOCATION] Previous Messages Context (Count: {len(messages)}) ================")
        for idx, msg in enumerate(messages):
            role_map = {"human": "USER", "ai": "ASSISTANT", "tool": "TOOL", "system": "SYSTEM"}
            role = role_map.get(msg.type, msg.type.upper())
            details = []
            if getattr(msg, "id", None):
                details.append(f"ID: {msg.id}")
            if msg.type == "ai" and getattr(msg, "tool_calls", None):
                details.append(f"ToolCalls: {[tc['name'] for tc in msg.tool_calls]}")
            elif msg.type == "tool":
                details.append(f"ToolName: {getattr(msg, 'name', '')}")
                details.append(f"ToolCallID: {getattr(msg, 'tool_call_id', '')}")
            
            details_str = f" ({', '.join(details)})" if details else ""
            content_str = msg.content
            if not isinstance(content_str, str):
                content_str = str(content_str)
            snippet = content_str[:120].replace('\n', ' ')
            if len(content_str) > 120:
                snippet += "..."
                
            print(f"  [{idx}] {role}: \"{snippet}\"{details_str}")
        print("=========================================================================================\n")

        # 1. Identify the user's latest prompt (iterating backwards from the end)
        user_prompt = ""
        latest_human_idx = -1
        for idx in range(len(messages) - 1, -1, -1):
            m = messages[idx]
            if isinstance(m, HumanMessage):
                latest_human_idx = idx
                if isinstance(m.content, str):
                    user_prompt = m.content.lower()
                elif isinstance(m.content, list) and len(m.content) > 0:
                    user_prompt = m.content[0].get("text", "").lower()
                break

        # 2. Extract conversation history details starting from the latest human message
        thought_count = 0
        system_tool_responses = 0
        
        if latest_human_idx != -1:
            for m in messages[latest_human_idx + 1:]:
                if isinstance(m, AIMessage) and m.tool_calls:
                    for tc in m.tool_calls:
                        if tc["name"] == "think":
                            thought_count += 1
                elif isinstance(m, ToolMessage) and m.name != "think":
                    system_tool_responses += 1

        import re
        user_prompt_original = messages[0].content if latest_human_idx == -1 else messages[latest_human_idx].content
        
        table_match = re.search(r'add table:(.*)', user_prompt_original, re.IGNORECASE)
        para_match = re.search(r'add paragraph:(.*)', user_prompt_original, re.IGNORECASE)

        # 3. Route decision trees based on the static mock_routes schemas
        invoc_id = uuid.uuid4().hex[:8]
        if any(kw in user_prompt for kw in ["multi", "multiple", "advanced", "complex", "step"]):
            ai_message = self._handle_multistep_route(thought_count, system_tool_responses, invoc_id)
        elif table_match or any(kw in user_prompt for kw in ["table", "grid", "rows", "cols", "columns"]):
            style_name = table_match.group(1).strip() if table_match else None
            ai_message = self._handle_add_table_route(thought_count, system_tool_responses, invoc_id, style_name)
        elif para_match or any(kw in user_prompt for kw in ["document", "edit", "write", "summary", "paragraph", "append"]):
            style_name = para_match.group(1).strip() if para_match else None
            ai_message = self._handle_edit_document_route(thought_count, system_tool_responses, invoc_id, style_name)
        elif any(kw in user_prompt for kw in ["search", "kb", "knowledge", "help", "guide", "info", "apcot"]):
            ai_message = self._handle_search_route(thought_count, system_tool_responses, messages[0].content, invoc_id)
        elif any(kw in user_prompt for kw in ["hello", "hi"]):
            ai_message = self._handle_greeting_route(thought_count, invoc_id)
        else:
            ai_message = self._handle_default_route(thought_count, messages[0].content, invoc_id)

        chat_generation = ChatGeneration(message=ai_message)
        return ChatResult(generations=[chat_generation])

    def _handle_multistep_route(self, thought_count: int, system_tool_responses: int, invoc_id: str) -> AIMessage:
        route = CONVERSATION_ROUTES["multi_step"]
        
        # Phase 1: 3 Consecutive Thought Steps
        if thought_count < 3:
            thought_text = route["thoughts_phase_1"][thought_count]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_p1_{thought_count}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
            
        # Tool Call 1: search_kb
        elif thought_count == 3 and system_tool_responses == 0:
            tc = route["tool_call_1"]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": tc["name"],
                    "args": tc["args"],
                    "id": f"tc_kb_search_1_{invoc_id}",
                    "type": "tool_call"
                }]
            )
            
        # Phase 2: 3 Consecutive Thought Steps after Tool 1
        elif system_tool_responses == 1 and thought_count < 6:
            idx = thought_count - 3
            thought_text = route["thoughts_phase_2"][idx]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_p2_{idx}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
            
        # Tool Call 2: check_entitlements
        elif system_tool_responses == 1 and thought_count == 6:
            tc = route["tool_call_2"]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": tc["name"],
                    "args": {"resource": tc["args"]["resource"]},
                    "id": f"tc_entitlements_check_2_{invoc_id}",
                    "type": "tool_call"
                }]
            )
            
        # Phase 3: 2 Consecutive Thought Steps after Tool 2
        elif system_tool_responses == 2 and thought_count < 8:
            idx = thought_count - 6
            thought_text = route["thoughts_phase_3"][idx]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_p3_{idx}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
            
        # Final Consolidated Response
        else:
            final_text = route["response_template"].format(
                fullname=self.fullname,
                user_id=self.user_id
            )
            return AIMessage(content=final_text)

    def _handle_edit_document_route(self, thought_count: int, system_tool_responses: int, invoc_id: str, style_name: str = None) -> AIMessage:
        route = CONVERSATION_ROUTES["edit_document"]
        para_ids = get_para_ids()
        p0 = para_ids[0] if len(para_ids) > 0 else "placeholder_id"
        p1 = para_ids[1] if len(para_ids) > 1 else p0
        
        # Phase 1: 2 thoughts
        if thought_count < 2:
            thought_text = route["thoughts"][thought_count]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_doc_{thought_count}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Tool Call 1: read_document
        elif thought_count == 2 and system_tool_responses == 0:
            tc = route["tool_call_1"]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": tc["name"],
                    "args": tc["args"],
                    "id": f"tc_read_doc_1_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Phase 2: 3 thoughts
        elif system_tool_responses == 1 and thought_count < 5:
            idx = thought_count - 2
            thought_text = route["thoughts_phase_2"][idx]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_doc_p2_{idx}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Tool Call 2: suggest_change
        elif system_tool_responses == 1 and thought_count == 5:
            tc = route["tool_call_2"]
            args = dict(tc["args"])
            args["paraId"] = p0
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": tc["name"],
                    "args": args,
                    "id": f"tc_suggest_change_2_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Phase 3: 2 thoughts
        elif system_tool_responses == 2 and thought_count < 7:
            idx = thought_count - 5
            thought_text = route["thoughts_phase_3"][idx]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_doc_p3_{idx}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Tool Call 3: add_comment
        elif system_tool_responses == 2 and thought_count == 7:
            tc = route["tool_call_3"]
            args = dict(tc["args"])
            args["paraId"] = p0
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": tc["name"],
                    "args": args,
                    "id": f"tc_add_comment_3_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Phase 4: 2 thoughts
        elif system_tool_responses == 3 and thought_count < 9:
            idx = thought_count - 7
            thought_text = route["thoughts_phase_4"][idx]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_doc_p4_{idx}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Tool Call 4: toggle_bullet_list
        elif system_tool_responses == 3 and thought_count == 9:
            tc = route["tool_call_4"]
            args = dict(tc["args"])
            args["paraId"] = p1
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": tc["name"],
                    "args": args,
                    "id": f"tc_toggle_bullet_4_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Phase 5: 2 thoughts
        elif system_tool_responses == 4 and thought_count < 11:
            idx = thought_count - 9
            thought_text = route["thoughts_phase_5"][idx]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_doc_p5_{idx}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Tool Call 5: append_paragraph
        elif system_tool_responses == 4 and thought_count == 11:
            tc = route["tool_call_5"]
            args = dict(tc["args"])
            if style_name:
                args["styleId"] = style_name
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": tc["name"],
                    "args": args,
                    "id": f"tc_append_para_5_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Phase 6: 2 thoughts
        elif system_tool_responses == 5 and thought_count < 13:
            idx = thought_count - 11
            thought_text = route["thoughts_phase_6"][idx]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_doc_p6_{idx}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Final response
        else:
            return AIMessage(content=route["response"])

    def _handle_add_table_route(self, thought_count: int, system_tool_responses: int, invoc_id: str, style_name: str = None) -> AIMessage:
        route = CONVERSATION_ROUTES["add_table"]
        para_ids = get_para_ids()
        p0 = para_ids[-1] if len(para_ids) > 0 else "placeholder_id"

        # Phase 1: 2 thoughts
        if thought_count < 2:
            thought_text = route["thoughts"][thought_count]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_tbl_{thought_count}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Tool Call 1: read_document
        elif thought_count == 2 and system_tool_responses == 0:
            tc = route["tool_call_1"]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": tc["name"],
                    "args": tc["args"],
                    "id": f"tc_read_doc_tbl_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Phase 2: 3 thoughts
        elif system_tool_responses == 1 and thought_count < 5:
            idx = thought_count - 2
            thought_text = route["thoughts_phase_2"][idx]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_tbl_p2_{idx}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Tool Call 2: insert_table
        elif system_tool_responses == 1 and thought_count == 5:
            tc = route["tool_call_2"]
            args = dict(tc["args"])
            args["paraId"] = p0
            if style_name:
                args["styleId"] = style_name
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": tc["name"],
                    "args": args,
                    "id": f"tc_insert_table_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Phase 3: 2 thoughts
        elif system_tool_responses == 2 and thought_count < 7:
            idx = thought_count - 5
            thought_text = route["thoughts_phase_3"][idx]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_tbl_p3_{idx}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Tool Call 3: add_table_row
        elif system_tool_responses == 2 and thought_count == 7:
            tc = route["tool_call_3"]
            args = dict(tc["args"])
            args["paraId"] = p0
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": tc["name"],
                    "args": args,
                    "id": f"tc_add_table_row_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Phase 4: 2 thoughts
        elif system_tool_responses == 3 and thought_count < 9:
            idx = thought_count - 7
            thought_text = route["thoughts_phase_4"][idx]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_tbl_p4_{idx}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        # Final response
        else:
            return AIMessage(content=route["response"])

    def _handle_search_route(self, thought_count: int, system_tool_responses: int, original_prompt: str, invoc_id: str) -> AIMessage:
        route = CONVERSATION_ROUTES["search"]
        
        # 3 Consecutive Thought Steps
        if thought_count < 3:
            thought_text = route["thoughts"][thought_count]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_search_{thought_count}_{invoc_id}",
                    "type": "tool_call"
                }]
            )
            
        # Tool Call: search_kb
        elif thought_count == 3 and system_tool_responses == 0:
            tc = route["tool_call"]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": tc["name"],
                    "args": {"query": original_prompt},
                    "id": f"tc_kb_search_1_{invoc_id}",
                    "type": "tool_call"
                }]
            )
            
        # Final text response
        else:
            return AIMessage(content=route["response"])
 
    def _handle_greeting_route(self, thought_count: int, invoc_id: str) -> AIMessage:
        route = CONVERSATION_ROUTES["greeting"]
        
        # 1 Thought Step
        if thought_count < 1:
            thought_text = route["thoughts"][0]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_greet_0_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        else:
            return AIMessage(content=route["response"])
 
    def _handle_default_route(self, thought_count: int, original_prompt: str, invoc_id: str) -> AIMessage:
        route = CONVERSATION_ROUTES["default"]
        
        # 1 Thought Step
        if thought_count < 1:
            thought_text = route["thoughts"][0]
            return AIMessage(
                content="",
                tool_calls=[{
                    "name": "think",
                    "args": {"thought": thought_text},
                    "id": f"tc_think_default_0_{invoc_id}",
                    "type": "tool_call"
                }]
            )
        else:
            final_text = route["response_template"].format(input_text=original_prompt)
            return AIMessage(content=final_text)
