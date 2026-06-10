# backend/app/services/document_service.py

import os
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Optional, List
import contextvars
from datetime import datetime
import random

# Thread-safe context variable for thread_id propagation
current_thread_id_var = contextvars.ContextVar("current_thread_id")

# Scoped storage directory
STORAGE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "storage", "documents")
os.makedirs(STORAGE_DIR, exist_ok=True)

def get_document_path(thread_id: str) -> str:
    return os.path.join(STORAGE_DIR, f"thread_{thread_id}.docx")

def enable_track_changes(doc):
    """Ensures OOXML Tracked Changes revisions setting is enabled on the document."""
    settings = doc.settings.element
    track_revisions = settings.find(qn('w:trackRevisions'))
    if track_revisions is None:
        track_revisions = OxmlElement('w:trackRevisions')
        settings.append(track_revisions)

import shutil

TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "assets", "report-template.docx")

def ensure_document_exists(thread_id: str) -> str:
    """Gets the document path. Creates a clean Word template if it does not exist."""
    path = get_document_path(thread_id)
    if not os.path.exists(path):
        if os.path.exists(TEMPLATE_PATH):
            shutil.copy(TEMPLATE_PATH, path)
            # Removed python-docx track changes injection because doc.save() 
            # drops template styles.xml data.
        else:
            doc = docx.Document()
            doc.add_heading("APCOT Chat Document Workspace", level=0)
            doc.add_paragraph("Welcome to your APCOT Chat Workspace Document. This document serves as the Single Source of Truth for this session.")
            doc.add_paragraph("You can edit this document here. Edits from both you (the user) and the APCOT Assistant will be tracked natively.")
            enable_track_changes(doc)
            doc.save(path)
    return path

def get_document_bytes(thread_id: str) -> bytes:
    """Fetches the document for a thread as raw bytes."""
    path = ensure_document_exists(thread_id)
    with open(path, "rb") as f:
        return f.read()

def save_document_bytes(thread_id: str, data: bytes):
    """Writes the document raw bytes to disk and enforces track revisions settings."""
    path = get_document_path(thread_id)
    with open(path, "wb") as f:
        f.write(data)
    # Removed python-docx save step because it aggressively drops advanced styles
    # from the styles.xml that the browser-based editor successfully generated.
    # The browser editor handles tracking changes natively.
    # try:
    #     doc = docx.Document(path)
    #     enable_track_changes(doc)
    #     doc.save(path)
    # except Exception as e:
    #     print("Error enforcing revisions on saved bytes:", e)

def apply_agent_edit(thread_id: str, action: str, text: str, paragraph_index: Optional[int] = None) -> str:
    """Applies programmatically-driven changes from the AI, wrapping them in native <w:ins> tracked elements."""
    path = ensure_document_exists(thread_id)
    try:
        doc = docx.Document(path)
        enable_track_changes(doc)

        def add_tracked_run(paragraph, text_content):
            """Appends a new run to paragraph, wrapping it in an OOXML tracked insertion tag."""
            p_element = paragraph._element
            ins_element = OxmlElement('w:ins')
            ins_element.set(qn('w:author'), 'APCOT Assistant')
            ins_element.set(qn('w:date'), datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ'))
            ins_element.set(qn('w:id'), str(random.randint(1000, 9999)))
            
            r_element = OxmlElement('w:r')
            t_element = OxmlElement('w:t')
            t_element.text = text_content
            r_element.append(t_element)
            ins_element.append(r_element)
            p_element.append(ins_element)

        if action == "append":
            p = doc.add_paragraph()
            add_tracked_run(p, text)
        elif action == "clear":
            # Clear all current paragraphs
            for _ in range(len(doc.paragraphs)):
                p_element = doc.paragraphs[0]._element
                p_element.getparent().remove(p_element)
            p = doc.add_paragraph()
            add_tracked_run(p, text)
        elif action == "replace":
            if paragraph_index is None:
                return "Error: paragraph_index is required for replace action."
            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return f"Error: paragraph_index {paragraph_index} is out of bounds (0 to {len(doc.paragraphs)-1})."
            p = doc.paragraphs[paragraph_index]
            p.text = "" # Clear runs
            add_tracked_run(p, text)
        elif action == "insert":
            if paragraph_index is None:
                return "Error: paragraph_index is required for insert action."
            if paragraph_index < 0 or paragraph_index > len(doc.paragraphs):
                return f"Error: paragraph_index {paragraph_index} is out of bounds (0 to {len(doc.paragraphs)})."
            
            if paragraph_index == len(doc.paragraphs):
                p = doc.add_paragraph()
                add_tracked_run(p, text)
            else:
                p_target = doc.paragraphs[paragraph_index]
                new_p = doc.add_paragraph()
                add_tracked_run(new_p, text)
                p_target._element.addprevious(new_p._element)
        elif action == "table":
            import json
            try:
                table_data = json.loads(text)
                if not isinstance(table_data, list) or not all(isinstance(row, list) for row in table_data):
                    return "Error: For table action, text must be a JSON-encoded 2D list of strings."
                
                rows = len(table_data)
                cols = len(table_data[0]) if rows > 0 else 0
                if rows == 0 or cols == 0:
                    return "Error: Table must have at least 1 row and 1 column."
                
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                for r_idx, row in enumerate(table_data):
                    for c_idx, val in enumerate(row):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = "" # Clear default
                        p = cell.paragraphs[0]
                        add_tracked_run(p, str(val))
            except Exception as e:
                return f"Error parsing table JSON data: {str(e)}"
        else:
            return f"Error: Unknown action '{action}'."

        doc.save(path)
        return f"Successfully updated document: {action} action completed."
    except Exception as e:
        return f"Error modifying document: {str(e)}"
